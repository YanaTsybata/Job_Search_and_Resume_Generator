from docx import Document

class ResumeGenerator():
    def __init__(self):
        # Create a new empty Word document
        self.document = Document()

    def create_basic_template(self):
        # Add "Resume" as a title at the beginning of the document
        self.document.add_heading('Resume', 0)

        #Add personal info
    def add_personal_info(self, name, email, phone):
        self.document.add_paragraph(name)
        self.document.add_paragraph(f"Email: {email}")
        self.document.add_paragraph(f"Phone: {phone}")

    def add_education(self, education_list):
        # Add "Education" section
        self.document.add_heading('Education', 1)
        #For each education, add a paragraph with information
        for edu in education_list:
            p = self.document.add_paragraph()
            p.add_run(f"{edu['degree']} in {edu['field']}").bold = True
            p.add_run(f"\n{edu['institution']}, {edu['year']}")

    def add_experience(self, experience_list):
        #Add "Work Experience" section
        self.document.add_heading('Work Experience', 1)
        #For each job, add a paragraph with information
        for exp in experience_list:
            p = self.document.add_paragraph()
            p.add_run(f"{exp['position']} at {exp['company']}").bold = True
            p.add_run(f"\n{exp['start_date']} - {exp['end_date']}")
            p.add_run(f"\n{exp['description']}")

    def add_skills(self, skills):
        # Add "Skills" section
        self.document.add_heading('Skills', 1)
        # Add all skills in one line and separated by commas
        skill_para = self.document.add_paragraph()
        skill_para.add_run(', '.join(skills))

    def customize_for_job(self, job_description):
        # Add a section describing how skills match the specific job
        self.document.add_heading('Job-Specific Qualifications', 1)
        self.document.add_paragraph(job_description)

    def save_resume(self, filename):
        # save resume document to a file
        self.document.save(filename)
        print(f"Resume saved as: {filename}")  #Add this line for debugging

    def create_resume(self, user_info, jobs_details):
        self.create_basic_template()  #Create the basic template
        self.add_personal_info(user_info['name'], user_info['email'], user_info['phone']) #Add user's personal information
        self.add_education(user_info['education']) #Add education section
        self.add_experience(user_info['experience']) #Add work experience section
        self.add_skills(user_info['skills']) #Add skills section

        self.document.add_page_break() #Add a page break before job
        #Add a section for "relevant job opportunities"
        self.document.add_heading('Relevant Job Opportunities', level=1)
        self.document.add_paragraph('Below are the job opportunities that match your profile:')

        #Add details for each job opportunity
        for job in jobs_details:
            self.document.add_heading(job.get('title', 'Job Title Not Available'), level=2)
            self.document.add_paragraph(f"Company: {job.get('company', 'Not Available')}")
            self.document.add_paragraph(f"Location: {job.get('location', 'Not Available')}")
            self.document.add_paragraph("Job Description:")
            self.document.add_paragraph(job.get('description', 'Description Not Available'))
            self.document.add_paragraph("---")  #Separator between job listings

        #Generate filename and save the resume
        filename = f"resume_and_jobs_{user_info['name'].replace(' ', '_')}.docx"
        self.save_resume(filename)
        return filename


#example
if __name__ == "__main__":
    user_info = {
        'name': 'Example Name',
        'email': 'exampleemail@example.com',
        'phone': '123-456-7890',
        'education': [
            {'degree': 'Bachelor', 'field': 'Computer Science', 'institution': 'XYZ University', 'year': '2021'}
        ],
        'experience': [
            {
                'position': 'Software Developer',
                'company': 'Tech Corp',
                'start_date': 'Jan 2021',
                'end_date': 'Present',
                'description': 'Developed web applications using Python and Django.'
            }
        ],
        'skills': ['Python', 'Django', 'JavaScript', 'SQL']
    }

    job_details = {
        'description': 'Looking for a Python developer with experience in web development.'
    }

    #Create and save the resume
    generator = ResumeGenerator()
    resume_file = generator.create_resume(user_info, job_details)
    print(f"Resume created: {resume_file}")
